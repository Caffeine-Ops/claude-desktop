"""docx_to_pdf 的行为契约。

重点全在「没装 LibreOffice 时会发生什么」——那是绝大多数用户的处境。
"""
import subprocess
import sys
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfReader

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))
import docx_to_pdf  # noqa: E402


def _make_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("季度汇报", level=1)
    doc.add_paragraph("这是一段中文正文，用来验证字体没有变成方块。")
    doc.save(str(path))


def test_refuses_textonly_without_explicit_flag(tmp_path, monkeypatch, capsys):
    monkeypatch.setattr(docx_to_pdf, "find_soffice", lambda: None)
    src = tmp_path / "a.docx"
    _make_docx(src)
    dst = tmp_path / "a.pdf"

    with pytest.raises(SystemExit) as e:
        docx_to_pdf.convert(src, dst, allow_textonly=False)

    assert e.value.code != 0
    err = capsys.readouterr().err
    # 报错必须同时说清：为什么不能转、装什么能解决、怎么强行继续
    assert "LibreOffice" in err
    assert "--allow-textonly" in err
    assert "排版" in err
    # 门禁的承诺：拒绝时不产出任何文件
    assert not dst.exists()


def test_textonly_path_produces_pdf_when_allowed(tmp_path, monkeypatch):
    monkeypatch.setattr(docx_to_pdf, "find_soffice", lambda: None)
    src = tmp_path / "b.docx"
    _make_docx(src)
    dst = tmp_path / "b.pdf"

    mode = docx_to_pdf.convert(src, dst, allow_textonly=True)

    assert mode == "textonly"
    assert dst.is_file()
    assert len(PdfReader(str(dst)).pages) >= 1


def test_textonly_refuses_when_no_cjk_font(tmp_path, monkeypatch, capsys):
    monkeypatch.setattr(docx_to_pdf, "find_soffice", lambda: None)
    monkeypatch.setattr(docx_to_pdf, "find_cjk_font", lambda: None)
    src = tmp_path / "c.docx"
    _make_docx(src)

    with pytest.raises(SystemExit):
        docx_to_pdf.convert(src, tmp_path / "c.pdf", allow_textonly=True)

    # 没有中文字体时输出的是满纸方块，不如不给
    assert "中文字体" in capsys.readouterr().err


def test_find_cjk_font_returns_existing_file_or_none():
    font = docx_to_pdf.find_cjk_font()
    assert font is None or font.is_file()


def test_soffice_failure_gives_chinese_message_not_traceback(tmp_path, monkeypatch, capsys):
    """评审后加固：soffice 崩溃/超时不能让 CalledProcessError 堆栈冒泡出去。

    最容易触发这条的场景是用户桌面正开着 LibreOffice，无头模式抢不到配置锁——
    这恰好是本技能默认主路径（装了 LibreOffice 就走这条），所以必须给中文提示。
    """
    monkeypatch.setattr(docx_to_pdf, "find_soffice", lambda: "/usr/bin/fake-soffice")

    def _boom(*args, **kwargs):
        raise subprocess.CalledProcessError(
            1, ["fake-soffice"], output=b"", stderr=b"some LibreOffice internal error\nsecond line"
        )

    monkeypatch.setattr(docx_to_pdf.subprocess, "run", _boom)
    src = tmp_path / "a.docx"
    _make_docx(src)
    dst = tmp_path / "a.pdf"

    with pytest.raises(SystemExit) as e:
        docx_to_pdf.convert(src, dst, allow_textonly=False)

    assert e.value.code != 0
    err = capsys.readouterr().err
    assert "LibreOffice" in err
    assert "关掉" in err  # 提醒"如果正开着，先关掉再试"
    assert not dst.exists()


def test_soffice_timeout_gives_chinese_message(tmp_path, monkeypatch, capsys):
    monkeypatch.setattr(docx_to_pdf, "find_soffice", lambda: "/usr/bin/fake-soffice")

    def _boom(*args, **kwargs):
        raise subprocess.TimeoutExpired(["fake-soffice"], 300)

    monkeypatch.setattr(docx_to_pdf.subprocess, "run", _boom)
    src = tmp_path / "a.docx"
    _make_docx(src)
    dst = tmp_path / "a.pdf"

    with pytest.raises(SystemExit):
        docx_to_pdf.convert(src, dst, allow_textonly=False)

    err = capsys.readouterr().err
    assert "超时" in err or "卡住" in err


def test_soffice_silent_no_op_gives_chinese_message(tmp_path, monkeypatch, capsys):
    """复审实测：soffice 还有一类「假成功」——退出码 0、stderr 干净，却一个
    PDF 都没写出来（源文件扩展名是 .docx 但内容不是 Word 文档最常见）。
    上面那个 try 包不住它（它包的是"调用失败"），原来会在紧接着的改名那步抛
    FileNotFoundError 英文堆栈。
    """
    monkeypatch.setattr(docx_to_pdf, "find_soffice", lambda: "/usr/bin/fake-soffice")

    class _SilentSuccess:
        returncode = 0
        stdout = b""
        stderr = b""

    # 装成「调用成功了」但不产出任何文件
    monkeypatch.setattr(docx_to_pdf.subprocess, "run", lambda *a, **k: _SilentSuccess())
    src = tmp_path / "weird.docx"
    src.write_bytes(b"GIF89a")  # 扩展名叫 docx，内容其实是别的格式
    dst = tmp_path / "weird.pdf"

    with pytest.raises(SystemExit) as e:
        docx_to_pdf.convert(src, dst, allow_textonly=False)

    assert e.value.code != 0
    err = capsys.readouterr().err
    assert "没有生成 PDF" in err
    assert "Traceback" not in err
    assert not dst.exists()


def test_main_wraps_unexpected_error_in_chinese(tmp_path, monkeypatch, capsys):
    """main() 的兜底层：任何未预期异常都要变成中文，不能漏英文堆栈。"""
    src = tmp_path / "a.docx"
    _make_docx(src)

    def _boom(*a, **k):
        raise OSError(30, "Read-only file system")

    monkeypatch.setattr(docx_to_pdf, "convert", _boom)

    with pytest.raises(SystemExit) as e:
        docx_to_pdf.main([str(src), "-o", str(tmp_path / "x.pdf")])

    assert e.value.code != 0
    err = capsys.readouterr().err
    assert "处理过程中出错" in err
    assert "Traceback" not in err


def test_doc_legacy_format_gives_chinese_message_on_textonly_path(tmp_path, monkeypatch, capsys):
    """评审后加固：.doc（旧二进制格式）走纯文字兜底时，python-docx 打不开会抛
    PackageNotFoundError，不能让它冒泡成英文堆栈——必须提示用户另存为 .docx。
    """
    monkeypatch.setattr(docx_to_pdf, "find_soffice", lambda: None)
    src = tmp_path / "legacy.doc"
    # 假装是一份旧版二进制 .doc：随便写点非 zip 内容，python-docx 打不开
    src.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1not a real doc file")
    dst = tmp_path / "legacy.pdf"

    with pytest.raises(SystemExit):
        docx_to_pdf.convert(src, dst, allow_textonly=True)

    err = capsys.readouterr().err
    assert "另存为" in err
    assert not dst.exists()


def test_empty_paragraph_document_refuses_textonly(tmp_path, monkeypatch, capsys):
    # 空段落文档（只有换行没有文字）应该被拒绝，不能静默产出空白 PDF
    monkeypatch.setattr(docx_to_pdf, "find_soffice", lambda: None)
    src = tmp_path / "d.docx"
    doc = Document()
    doc.add_paragraph("")  # 纯空段落
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.save(str(src))
    dst = tmp_path / "d.pdf"

    with pytest.raises(SystemExit):
        docx_to_pdf.convert(src, dst, allow_textonly=True)

    # 拒绝时不产出文件
    assert not dst.exists()
    assert "没有可提取的文字" in capsys.readouterr().err


def test_find_soffice_falls_back_to_default_install_paths(tmp_path, monkeypatch):
    """Windows CI 验证的后续发现（2026-08-14）：LibreOffice 的 Windows 安装器
    默认不写 PATH，which 永远找不到——原实现只点名了 macOS 默认路径，
    Windows 用户装了 LibreOffice 也会被当成没装、永远走不到保排版路径。
    改成扫 _SOFFICE_DEFAULTS 候选表（mac + Windows 两个 Program Files），
    本测试用假候选表钉住「which 落空 → 扫默认路径」这条链路本身。"""
    fake = tmp_path / "soffice.exe"
    fake.write_bytes(b"")
    monkeypatch.setattr(docx_to_pdf.shutil, "which", lambda name: None)
    monkeypatch.setattr(docx_to_pdf, "_SOFFICE_DEFAULTS", [str(tmp_path / "不存在"), str(fake)])
    assert docx_to_pdf.find_soffice() == str(fake)


def test_find_soffice_returns_none_when_nothing_found(tmp_path, monkeypatch):
    monkeypatch.setattr(docx_to_pdf.shutil, "which", lambda name: None)
    monkeypatch.setattr(docx_to_pdf, "_SOFFICE_DEFAULTS", [str(tmp_path / "不存在")])
    assert docx_to_pdf.find_soffice() is None


def test_soffice_default_candidates_cover_windows_and_mac():
    """反遗漏断言：候选表必须同时点名 macOS 与 Windows 的默认安装路径——
    这正是本次要修的缺口，谁把 Windows 条目删了这条会当场红。"""
    joined = "\n".join(docx_to_pdf._SOFFICE_DEFAULTS)
    assert "/Applications/LibreOffice.app" in joined
    assert "Program Files\\LibreOffice" in joined or "Program Files/LibreOffice" in joined
