"""doc_text 的行为契约。

最要紧的是体检报告，尤其 scanned 判定：长文档提炼最大的风险不是总结得不好，
是模型只读了前面一小截就开始总结、而且它不会告诉你。取料时就把「多少页、
多少字、哪几页是扫描的」摊在台面上，agent 才有依据决定分不分块、走不走 OCR。
"""
import json
import subprocess
import sys
from pathlib import Path

import pytest
from PIL import Image

SCRIPTS = Path(__file__).resolve().parents[1] / "scripts"
sys.path.insert(0, str(SCRIPTS))
import doc_text  # noqa: E402


def _text_pdf(path: Path, pages: int = 2) -> Path:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    c = canvas.Canvas(str(path), pagesize=A4)
    for i in range(pages):
        c.drawString(72, 720, f"This is page {i + 1} with enough text to look real. " * 4)
        c.showPage()
    c.save()
    return path


def _image_only_pdf(path: Path, tmp_path: Path) -> Path:
    """造一份「扫描件」：整页只有一张图，没有任何文字层。"""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    img = tmp_path / "blank.png"
    Image.new("RGB", (800, 1100), (240, 240, 240)).save(img)
    c = canvas.Canvas(str(path), pagesize=A4)
    c.drawImage(str(img), 0, 0, width=A4[0], height=A4[1])
    c.showPage()
    c.save()
    return path


def _docx(path: Path) -> Path:
    from docx import Document
    d = Document()
    d.add_paragraph("第一段内容")
    d.add_paragraph("第二段内容")
    d.save(str(path))
    return path


def test_text_pdf_is_not_scanned(tmp_path):
    units, kind = doc_text.extract(_text_pdf(tmp_path / "a.pdf"))
    report = doc_text.checkup(units, kind)
    assert kind == "pdf"
    assert report["units"] == 2
    assert report["scanned"] is False
    assert report["scanned_units"] == []


def test_image_only_pdf_is_flagged_scanned(tmp_path):
    units, kind = doc_text.extract(_image_only_pdf(tmp_path / "s.pdf", tmp_path))
    report = doc_text.checkup(units, kind)
    assert report["scanned"] is True
    assert report["scanned_units"] == [1]


def test_docx_units_are_paragraphs(tmp_path):
    units, kind = doc_text.extract(_docx(tmp_path / "a.docx"))
    assert kind == "docx"
    assert units == ["第一段内容", "第二段内容"]


def test_legacy_doc_is_refused_in_chinese(tmp_path):
    old = tmp_path / "old.doc"
    old.write_bytes(b"\xd0\xcf\x11\xe0")
    proc = subprocess.run(
        [sys.executable, str(SCRIPTS / "doc_text.py"), str(old), "--outdir", str(tmp_path)],
        capture_output=True, text=True,
    )
    assert proc.returncode != 0
    assert "另存为" in proc.stderr
    assert "Traceback" not in proc.stderr


def test_cli_writes_text_file_with_page_anchors(tmp_path):
    src = _text_pdf(tmp_path / "b.pdf", pages=3)
    proc = subprocess.run(
        [sys.executable, str(SCRIPTS / "doc_text.py"), str(src), "--outdir", str(tmp_path)],
        capture_output=True, text=True,
    )
    assert proc.returncode == 0, proc.stderr
    report = json.loads(proc.stdout)
    body = Path(report["text_file"]).read_text(encoding="utf-8")
    assert "[P1]" in body and "[P3]" in body


def test_cli_docx_uses_paragraph_anchors(tmp_path):
    src = _docx(tmp_path / "c.docx")
    proc = subprocess.run(
        [sys.executable, str(SCRIPTS / "doc_text.py"), str(src), "--outdir", str(tmp_path)],
        capture_output=True, text=True,
    )
    assert proc.returncode == 0, proc.stderr
    body = Path(json.loads(proc.stdout)["text_file"]).read_text(encoding="utf-8")
    assert "[§1] 第一段内容" in body


# --- 以下两条是本 PR 的补充要求：main() 的兜底防线不能只靠零散 try/except，
# 写盘调用（mkdir / write_text）失败必须转成中文错误，不能把裸 Traceback
# 甩给用户。用 monkeypatch 强制这两个调用抛异常，而不是真的填满磁盘或
# 改权限——那样在 CI 环境里既不稳定也不方便清理。

def test_mkdir_failure_gives_chinese_message_not_traceback(tmp_path, monkeypatch, capsys):
    src = _text_pdf(tmp_path / "a.pdf")

    def _boom(self, parents=True, exist_ok=True):
        raise OSError("模拟的目录创建失败")

    monkeypatch.setattr(Path, "mkdir", _boom)
    with pytest.raises(SystemExit) as exc:
        doc_text.main([str(src), "--outdir", str(tmp_path / "out")])
    assert exc.value.code != 0
    captured = capsys.readouterr()
    assert "[doc-convert] 错误：" in captured.err
    assert "Traceback" not in captured.err


def test_write_text_failure_gives_chinese_message_not_traceback(tmp_path, monkeypatch, capsys):
    src = _text_pdf(tmp_path / "a.pdf")

    def _boom(self, *args, **kwargs):
        raise OSError("模拟的磁盘写入失败")

    monkeypatch.setattr(Path, "write_text", _boom)
    with pytest.raises(SystemExit) as exc:
        doc_text.main([str(src), "--outdir", str(tmp_path)])
    assert exc.value.code != 0
    captured = capsys.readouterr()
    assert "[doc-convert] 错误：" in captured.err
    assert "Traceback" not in captured.err


# --- 补充：加密探测必须是类型判断，不是猜英文措辞。评审判定字符串匹配是
# Important 级缺陷——库升级改一句措辞，加密 PDF 就会滑进「文件可能已损坏」
# 的通用分支，答非所问。用 pypdf 现造一份真正带密码的 PDF，不能凭记忆假设
# 异常类型。同一处缺陷在 pdf_tables.py 的 _open 里也修了，两处判断逻辑一致。

def test_extract_pdf_reports_password_protection_in_chinese(tmp_path, capsys):
    """实测：pdfplumber.open() 对加密 PDF 抛出的是
    pdfplumber.utils.exceptions.PdfminerException，str(e) 是空的——原来的
    字符串匹配对这个包装类从来没生效过。真正管用的信号在 e.args[0]，那是
    被包了一层的原始异常 pdfminer.pdfdocument.PDFPasswordIncorrect，
    继承自 PDFEncryptionError。"""
    from pypdf import PdfReader, PdfWriter

    src = _text_pdf(tmp_path / "h.pdf", pages=1)
    reader = PdfReader(str(src))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt(user_password="secret")
    enc = tmp_path / "enc.pdf"
    with enc.open("wb") as f:
        writer.write(f)

    with pytest.raises(SystemExit):
        doc_text._extract_pdf(enc)
    err = capsys.readouterr().err
    assert "密码保护" in err


def _run_main(argv: list[str], capsys) -> dict:
    assert doc_text.main(argv) == 0
    return json.loads(capsys.readouterr().out)


def test_same_stem_different_formats_do_not_overwrite(tmp_path, capsys):
    """复审实测：产物名原来只取主干名（stem），`年报.pdf` 与 `年报.docx` 进同一个
    --outdir 会算出同一个路径，后跑的静默盖掉先跑的、exit 0 无任何提示。
    而 SKILL.md 的 A4 工作流恰恰教模型把多份输入都指到同一个 `取料/`。
    """
    from docx import Document

    pdf = _text_pdf(tmp_path / "年报.pdf")
    docx = tmp_path / "年报.docx"
    d = Document()
    d.add_paragraph("Word 文档里的正文，内容与那份 PDF 完全不同。")
    d.save(str(docx))
    outdir = tmp_path / "取料"

    r1 = _run_main([str(pdf), "--outdir", str(outdir)], capsys)
    r2 = _run_main([str(docx), "--outdir", str(outdir)], capsys)

    f1, f2 = Path(r1["text_file"]), Path(r2["text_file"])
    assert f1 != f2
    assert f1.is_file() and f2.is_file()
    # 两份取料内容必须各自完好，谁也没被谁盖掉
    assert "This is page 1" in f1.read_text(encoding="utf-8")
    assert "Word 文档里的正文" in f2.read_text(encoding="utf-8")


def test_rerunning_same_source_is_idempotent(tmp_path, capsys):
    """撞名护栏不能带来副作用：同一份源重跑要落回同一个文件，不能越跑越多。"""
    pdf = _text_pdf(tmp_path / "年报.pdf")
    outdir = tmp_path / "取料"

    r1 = _run_main([str(pdf), "--outdir", str(outdir)], capsys)
    r2 = _run_main([str(pdf), "--outdir", str(outdir)], capsys)

    assert r1["text_file"] == r2["text_file"]
    assert len(list(outdir.glob("*.text.txt"))) == 1


def test_atomic_write_success_leaves_no_temp_file(tmp_path):
    """挂账收口：结果文件必须原子落盘——成功后目录里只有目标文件，无 .part 残留。"""
    dst = tmp_path / "r.txt"
    doc_text._write_text_atomic(dst, '{"ok": 1}')
    assert dst.read_text(encoding="utf-8") == '{"ok": 1}'
    assert [p.name for p in tmp_path.iterdir()] == ["r.txt"]


def test_atomic_write_failure_leaves_no_partial_file(tmp_path, monkeypatch):
    """挂账收口：替换（os.replace）失败时不能留下半截临时文件，目标文件也不能
    出现——半截 JSON 比没有更糟，下游会拿着残缺数据继续走。"""
    def _boom(src, dst):
        raise OSError("simulated failure")

    monkeypatch.setattr(doc_text.os, "replace", _boom)
    dst = tmp_path / "r.txt"
    with pytest.raises(OSError):
        doc_text._write_text_atomic(dst, "{}")
    assert not dst.exists()
    assert list(tmp_path.iterdir()) == []


def test_zero_page_pdf_is_refused(tmp_path):
    """挂账收口：0 页的合法 PDF 原来会静默产出空取料文件 + exit 0，与「只有
    空段落的 docx 走拒绝」不对称。两条路径的纪律要一致：给不了任何有用产物
    就拒绝，不产出一份空文件让下游误以为「读完了、只是没内容」。"""
    from pypdf import PdfWriter

    src = tmp_path / "empty.pdf"
    with src.open("wb") as f:
        PdfWriter().write(f)
    outdir = tmp_path / "取料"
    proc = subprocess.run(
        [sys.executable, str(SCRIPTS / "doc_text.py"), str(src), "--outdir", str(outdir)],
        capture_output=True, text=True,
    )
    assert proc.returncode != 0
    assert proc.stderr.startswith("[doc-convert] 错误：")
    assert "0 页" in proc.stderr
    assert not outdir.exists() or list(outdir.glob("*.text.txt")) == []
