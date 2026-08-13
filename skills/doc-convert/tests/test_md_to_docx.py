"""md_to_docx 的行为契约。

只测「结构映射对不对」，不测排版细节——排版由 Word 默认样式决定，
断言它等于把 python-docx 的实现细节钉进测试里。
"""
import sys
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))
import md_to_docx  # noqa: E402


def test_headings_paragraphs_and_bullets(tmp_path):
    src = tmp_path / "a.md"
    src.write_text(
        "# 标题一\n\n正文一段。\n\n- 项目甲\n- 项目乙\n\n## 标题二\n",
        encoding="utf-8",
    )
    dst = tmp_path / "a.docx"

    md_to_docx.convert(src, dst)

    doc = Document(str(dst))
    texts = [p.text for p in doc.paragraphs]
    styles = [p.style.name for p in doc.paragraphs]
    assert texts[:5] == ["标题一", "正文一段。", "项目甲", "项目乙", "标题二"]
    assert styles[0] == "Heading 1"
    assert styles[1] == "Normal"
    assert styles[2] == "List Bullet"
    assert styles[4] == "Heading 2"


def test_bold_and_italic_become_runs(tmp_path):
    src = tmp_path / "b.md"
    src.write_text("这是**重点**和*强调*。\n", encoding="utf-8")
    dst = tmp_path / "b.docx"

    md_to_docx.convert(src, dst)

    doc = Document(str(dst))
    runs = doc.paragraphs[0].runs
    assert "".join(r.text for r in runs) == "这是重点和强调。"
    assert any(r.bold and r.text == "重点" for r in runs)
    assert any(r.italic and r.text == "强调" for r in runs)


def test_fenced_code_block_kept_verbatim(tmp_path):
    src = tmp_path / "c.md"
    src.write_text("```\nline1\nline2\n```\n", encoding="utf-8")
    dst = tmp_path / "c.docx"

    md_to_docx.convert(src, dst)

    doc = Document(str(dst))
    # 代码块逐行成段，且不被行内标记解析（`*` 之类原样保留）
    assert [p.text for p in doc.paragraphs[:2]] == ["line1", "line2"]


def test_missing_input_exits_with_message(tmp_path, capsys):
    with pytest.raises(SystemExit) as e:
        md_to_docx.main([str(tmp_path / "nope.md"), "-o", str(tmp_path / "x.docx")])
    assert e.value.code != 0
    assert "找不到输入文件" in capsys.readouterr().err


def test_bom_prefixed_markdown_still_recognizes_heading(tmp_path):
    """复审实测：Windows 记事本 / VS Code 存 UTF-8 时会在文件开头塞三个不可见
    字节（BOM）。原来用 encoding="utf-8" 读，BOM 会留在首行变成 `\\ufeff# 标题`，
    匹配不上标题规则、被当普通正文默默写进 Word——exit 0，用户拿到的文档少了
    一级标题却毫无提示。
    """
    src = tmp_path / "bom.md"
    src.write_bytes("﻿# 标题一\n\n正文一段。\n".encode("utf-8"))
    dst = tmp_path / "bom.docx"

    md_to_docx.convert(src, dst)

    doc = Document(str(dst))
    assert doc.paragraphs[0].text == "标题一"
    assert doc.paragraphs[0].style.name == "Heading 1"


def test_gbk_markdown_is_readable(tmp_path):
    """中文 Windows 上另存的 .md 常是 GBK，原来直接抛 UnicodeDecodeError 英文堆栈。"""
    src = tmp_path / "gbk.md"
    src.write_bytes("# 中文标题\n\n中文正文。\n".encode("gbk"))
    dst = tmp_path / "gbk.docx"

    md_to_docx.convert(src, dst)

    doc = Document(str(dst))
    assert doc.paragraphs[0].text == "中文标题"
    assert doc.paragraphs[1].text == "中文正文。"


def test_undecodable_bytes_give_chinese_message(tmp_path, capsys):
    """UTF-8 与 GB18030 都解不了时要给中文提示，不是英文堆栈。"""
    src = tmp_path / "bad.md"
    src.write_bytes(b"\xff\xfe\x00# x\x00")  # UTF-16 开头，两种候选编码都解不通

    with pytest.raises(SystemExit) as e:
        md_to_docx.convert(src, tmp_path / "bad.docx")

    assert e.value.code != 0
    err = capsys.readouterr().err
    assert "编码" in err
    assert "Traceback" not in err


def test_main_wraps_unexpected_error_in_chinese(tmp_path, monkeypatch, capsys):
    """main() 的兜底层：doc.save() 遇到只读目录抛的裸 OSError 不能漏出去。"""
    src = tmp_path / "a.md"
    src.write_text("# 标题\n", encoding="utf-8")

    def _boom(*a, **k):
        raise OSError(30, "Read-only file system")

    monkeypatch.setattr(md_to_docx, "convert", _boom)

    with pytest.raises(SystemExit) as e:
        md_to_docx.main([str(src), "-o", str(tmp_path / "a.docx")])

    assert e.value.code != 0
    err = capsys.readouterr().err
    assert "处理过程中出错" in err
    assert "Traceback" not in err


from docx.enum.text import WD_ALIGN_PARAGRAPH


def test_pipe_table_becomes_word_table(tmp_path):
    """表格是 A 类主力场景（发票/报表）最常见的结构，此前会塌成竖线文本。"""
    src = tmp_path / "t.md"
    src.write_text(
        "| 品名 | 金额 |\n|:---|---:|\n| 差旅 | 128.5 |\n| 餐饮 | 56 |\n",
        encoding="utf-8",
    )
    dst = tmp_path / "t.docx"
    md_to_docx.convert(src, dst)
    doc = Document(str(dst))
    assert len(doc.tables) == 1
    t = doc.tables[0]
    assert t.rows[0].cells[0].text == "品名"
    assert t.rows[2].cells[1].text == "56"
    # 表头加粗；右对齐列（---:）落到单元格段落上
    assert t.rows[0].cells[0].paragraphs[0].runs[0].bold is True
    assert t.rows[1].cells[1].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    # 竖线文本不应再出现在正文段落里
    assert all("|" not in p.text for p in doc.paragraphs)


def test_table_cell_inline_formatting_works(tmp_path):
    src = tmp_path / "u.md"
    src.write_text("| 项 |\n|---|\n| **重点** |\n", encoding="utf-8")
    dst = tmp_path / "u.docx"
    md_to_docx.convert(src, dst)
    cell_runs = Document(str(dst)).tables[0].rows[1].cells[0].paragraphs[0].runs
    assert cell_runs[0].text == "重点" and cell_runs[0].bold is True


def test_short_row_is_padded_with_empty_cells(tmp_path):
    """GFM 标准行为：短行补空，不丢信息、不拒绝。"""
    src = tmp_path / "v.md"
    src.write_text("| 甲 | 乙 |\n|---|---|\n| 只有一格 |\n", encoding="utf-8")
    dst = tmp_path / "v.docx"
    md_to_docx.convert(src, dst)
    t = Document(str(dst)).tables[0]
    assert t.rows[1].cells[0].text == "只有一格"
    assert t.rows[1].cells[1].text == ""


def test_long_row_is_refused_with_line_number(tmp_path, capsys):
    """纪律：比表头长的行如果截断就丢内容，宁可拒绝，并报出 md 行号。"""
    src = tmp_path / "w.md"
    src.write_text("| 甲 | 乙 |\n|---|---|\n| 1 | 2 | 3 |\n", encoding="utf-8")
    dst = tmp_path / "w.docx"
    with pytest.raises(SystemExit):
        md_to_docx.convert(src, dst)
    err = capsys.readouterr().err
    assert err.startswith("[doc-convert] 错误：")
    assert "第 3 行" in err and "列" in err
    assert not dst.exists()


def test_pipe_lines_without_separator_stay_plain_text(tmp_path):
    """没有分隔行就不是表格——维持旧行为当普通文本，不误伤正文里的竖线。"""
    src = tmp_path / "x.md"
    src.write_text("| 这行只是碰巧有竖线 |\n没有分隔行。\n", encoding="utf-8")
    dst = tmp_path / "x.docx"
    md_to_docx.convert(src, dst)
    doc = Document(str(dst))
    assert doc.tables == []
    assert "| 这行只是碰巧有竖线 |" in [p.text for p in doc.paragraphs]


def test_table_syntax_inside_code_fence_is_not_parsed(tmp_path):
    src = tmp_path / "y.md"
    src.write_text("```\n| a | b |\n|---|---|\n| 1 | 2 |\n```\n", encoding="utf-8")
    dst = tmp_path / "y.docx"
    md_to_docx.convert(src, dst)
    assert Document(str(dst)).tables == []
