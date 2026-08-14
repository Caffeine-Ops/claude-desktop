#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""doc_text.py — 统一取料 + 体检：PDF / Word(.docx) / txt / md → 带锚点纯文本。

体检报告才是这个脚本存在的主要理由。长文档提炼最大的风险不是总结得不好，
是模型只读了前面一小截就开始总结——而且它不会告诉你。所以取料这一步就得把
「这份文档多少页、多少字、哪几页根本没有文字层」摊在台面上，让 agent 据此
决定分不分块、要不要改走 OCR 路线，而不是闷头读完开头就下结论。

锚点是提取后自编的定位坐标（PDF 用页号 [P3]，其余用段号 [§12]），
文件本身没有。摘要里的结论要带出处，靠的就是它。
做法与 skills/tender-review/scripts/extract_text.py 的行号锚点同源。
"""
import argparse
import json
import os
import sys
from pathlib import Path

SCANNED_CHARS_PER_PAGE = 50  # 平均每页可提取字符低于此值 → 判定扫描件

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass


def _die(msg: str) -> None:
    print(f"[doc-convert] 错误：{msg}", file=sys.stderr)
    raise SystemExit(2)


def _write_text_atomic(dst: Path, body: str) -> None:
    """先写同目录临时文件、成功后原子改名，失败清掉临时文件再抛。

    挂账收口（PR #31 Task 5 deferred）：原来的 dst.write_text 直接往目标路径写，
    中途失败（磁盘满/进程被杀）会留下半截文件——下游拿着残缺 JSON 继续走，
    比报错更糟。临时文件必须与目标同目录：os.replace 跨文件系统不保证原子。
    """
    tmp = dst.with_name(dst.name + ".part")
    try:
        tmp.write_text(body, encoding="utf-8")
        os.replace(tmp, dst)
    except Exception:
        tmp.unlink(missing_ok=True)
        raise


def _extract_pdf(path: Path) -> list[str]:
    """打开并读取 PDF，加密/损坏都转成中文报错。

    评审加固（同 pdf_render.py 的 open_document、pdf_tables.py 的 _open）：
    加密探测要靠类型判断，不能靠猜英文措辞。2026-08-11 用 pypdf 现造一份
    真正带密码的 PDF 实测确认：pdfplumber.open() 对它抛出的是
    pdfplumber.utils.exceptions.PdfminerException，**str(e) 是空字符串**——
    原来那版 "password" in str(e).lower() 的字符串匹配对这个包装类从来没
    生效过，纯靠巧合没被任何测试戳穿。真正管用的信号在 e.args[0]：pdfplumber
    把被捕获的原始异常整个塞进 PdfminerException(e) 的构造参数，e.args[0]
    就是那个原始异常——pdfminer 里加密相关的问题都继承自
    pdfminer.pdfdocument.PDFEncryptionError，isinstance 判断稳。字符串匹配
    保留做兜底，不再是主判据。
    """
    import pdfplumber
    from pdfminer.pdfdocument import PDFEncryptionError
    try:
        with pdfplumber.open(str(path)) as pdf:
            return [(p.extract_text() or "") for p in pdf.pages]
    except Exception as e:
        inner = e.args[0] if e.args else None
        low = str(e).lower()
        if isinstance(inner, PDFEncryptionError) or "password" in low or "encrypt" in low:
            _die(f"PDF「{path.name}」被密码保护，本工具无法处理。请先用阅读器去掉密码再试。")
        _die(f"打开 PDF「{path.name}」失败，文件可能已损坏。请确认后重试。")
        return []  # 不可达，只为类型完整


def _extract_docx(path: Path) -> list[str]:
    """按文档顺序遍历段落和表格。

    python-docx 默认把 doc.paragraphs 和 doc.tables 分开返回，丢失原文顺序；
    这里手动遍历 body 的子元素，保持正文与表格的真实先后——同 tender-review
    的 extract_text.py。摘要引用位置时顺序错了，出处就是错的。
    """
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    try:
        doc = Document(str(path))
    except Exception:
        _die(f"打开 Word 文档「{path.name}」失败，文件可能已损坏。请确认后重试。")
    units: list[str] = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            txt = Paragraph(child, doc).text.strip()
            if txt:
                units.append(txt)
        elif child.tag == qn("w:tbl"):
            for row in Table(child, doc).rows:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                units.append(" | ".join(cells))
    return units


def _extract_plain(path: Path) -> list[str]:
    try:
        raw = path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        _die(f"读取「{path.name}」失败。请确认文件没有损坏。")
    return [ln.strip() for ln in raw.splitlines() if ln.strip()]


def extract(path: Path) -> tuple[list[str], str]:
    """返回 (分段文本, kind)。PDF 一项一页，其余一项一段。"""
    path = Path(path)
    if not path.is_file():
        _die(f"找不到文件「{path}」。请确认路径。")
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path), "pdf"
    if suffix == ".docx":
        return _extract_docx(path), "docx"
    if suffix == ".doc":
        _die("这是旧的 .doc 二进制格式，本工具打不开。请先在 Word 里另存为 .docx 或 PDF 再试。")
    if suffix in {".txt", ".md", ".markdown"}:
        return _extract_plain(path), "plain"
    _die(f"不支持的格式「{suffix}」。本能力只吃 PDF / .docx / .txt / .md。")
    return [], ""  # 不可达


def checkup(units: list[str], kind: str) -> dict:
    """体检报告。scanned 只对 PDF 有意义——其余格式没有「文字层」这回事。"""
    per_unit = [len(u) for u in units]
    total = sum(per_unit)
    scanned_units: list[int] = []
    scanned = False
    if kind == "pdf" and units:
        # 按页给明细而不是一刀切：混合型文档（前半电子版、后半扫描插页）
        # 很常见，只报一个 true/false 会让 agent 对整份文档做错决定。
        scanned_units = [i + 1 for i, n in enumerate(per_unit) if n < SCANNED_CHARS_PER_PAGE]
        scanned = (total / len(units)) < SCANNED_CHARS_PER_PAGE
    return {
        "kind": kind,
        "units": len(units),
        "chars": total,
        "chars_per_unit": per_unit,
        "scanned": scanned,
        "scanned_units": scanned_units,
    }


def _resolve_text_path(outdir: Path, src: Path, body: str) -> Path:
    """算出取料文件该叫什么，保证不会悄悄盖掉别人的产物。

    复审实测（同 img_prep.prepare_one 那条护栏的同族问题）：原来的名字是
    `src.stem + ".text.txt"`，**只取主干名、丢掉扩展名**。于是 `年报.pdf` 和
    `年报.docx` 进同一个 --outdir 会算出同一个路径，后跑的静默盖掉先跑的，
    exit 0 无任何提示。而 SKILL.md 的 A4 工作流恰恰教模型把多份输入都指到同一个
    `取料/`，这不是理论风险。

    两层解法：
      1. 名字带上扩展名（`年报.pdf.text.txt` / `年报.docx.text.txt`），把最常见
         的"同名不同格式"从根上拆开，而且同一份源文件重跑仍然落回同一个名字
         （幂等，不会每跑一次多出一个文件）。
      2. 仍然撞名时（两个不同目录下的同名同格式文件进同一个 outdir）比内容：
         内容一样说明就是上一次跑同一份源留下的，直接覆盖；内容不一样才换成
         `年报.pdf-2.text.txt` 让开——宁可多一个文件，也不让用户的取料结果
         凭空消失。
    """
    base = src.name  # 带扩展名，这是与旧版唯一的区别
    candidate = outdir / f"{base}.text.txt"
    n = 2
    while candidate.exists():
        try:
            if candidate.read_text(encoding="utf-8") == body:
                return candidate  # 上一次跑同一份源留下的，覆盖即幂等重跑
        except OSError:
            pass  # 读不了就当作"内容不同"，换名让开，不冒覆盖的险
        candidate = outdir / f"{base}-{n}.text.txt"
        n += 1
    return candidate


def render_anchored(units: list[str], kind: str) -> str:
    tag = "P" if kind == "pdf" else "§"
    return "\n\n".join(f"[{tag}{i}] {u}" for i, u in enumerate(units, start=1))


def main(argv: list[str] | None = None) -> int:
    try:
        ap = argparse.ArgumentParser(description="统一取料：文档 → 带锚点纯文本 + 体检报告")
        ap.add_argument("input", help="PDF / .docx / .txt / .md")
        ap.add_argument("--outdir", default=".", help="文本产物目录，默认当前目录")
        args = ap.parse_args(argv)

        src = Path(args.input)
        units, kind = extract(src)
        if not units:
            # 一个字都提不出来又不是扫描件判定能解释的，属于「给不了任何有用产物」。
            # PDF 侧：units 一项一页，空列表 = 0 页文件（扫描件每页仍占一项，不会
            # 走到这里）。原来这个分支只拦非 PDF，0 页 PDF 会静默产出空取料文件
            # + exit 0——与「只有空段落的 docx 被拒绝」不对称（挂账收口）。
            if kind == "pdf":
                _die(f"「{src.name}」是一份 0 页的 PDF，没有内容可提取。请确认文件是否正确。")
            _die(f"「{src.name}」里提不出任何文字。请确认文件内容是否正确。")

        report = checkup(units, kind)
        outdir = Path(args.outdir)
        try:
            outdir.mkdir(parents=True, exist_ok=True)
        except Exception:
            # 写盘调用单独兜底，不能让磁盘满/权限拒绝这类系统层异常
            # 裸露成 Traceback——同 img_prep.py 的纪律。
            _die(f"无法创建输出目录 {outdir}，请检查目录权限或磁盘空间。")

        body = render_anchored(units, kind)
        text_file = _resolve_text_path(outdir, src, body)
        try:
            _write_text_atomic(text_file, body)
        except Exception:
            _die(f"写入文本文件 {text_file} 失败，请检查目标目录权限或磁盘空间。")

        report["source"] = src.name
        report["text_file"] = str(text_file)
        print(json.dumps(report, ensure_ascii=False, indent=2))
        return 0
    except Exception as e:
        # 兜底：main() 必须有这一层，逐个函数自觉包 try 是不够的——任何未
        # 预期的异常（哪怕来自看起来无辜的一行代码）都要转成中文错误，
        # 不能让裸 Traceback 打到用户面前。这是本 PR 的全局约束，
        # 同 img_prep.py / md_to_docx.py 等脚本的纪律保持一致。
        _die(f"处理过程中出错：{type(e).__name__}: {e}")


if __name__ == "__main__":
    raise SystemExit(main())
