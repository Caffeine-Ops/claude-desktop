#!/usr/bin/env python3
"""Markdown → Word（.docx）。

**刻意只支持 Markdown 的常用子集**，不引 markdown/mistune 之类解析库：
用户在办公场景写的 md 就是标题 / 段落 / 列表 / 表格 / 粗斜体 / 代码块这几样，
解析全靠手写规则就够。真遇到复杂 md，正确做法是让模型先把它规整成这个子集，
而不是把解析器做厚。

支持：# ~ ###### 标题、空行分段、- / * 无序列表、1. 有序列表、
     **粗体**、*斜体*、`行内代码`、``` 围栏代码块、--- 分隔线、
     | 管道表格（表头加粗、:--- 对齐语法；2026-08-13 补，A 类场景刚需）。
不支持（会原样当普通文本输出，不报错）：引用块、图片、链接语法、
     表格里的合并单元格与 \\| 转义竖线。
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document

# 行内标记：粗体优先于斜体（`**x**` 必须先被吃掉，否则会被斜体规则劈成
# `*` + `*x*` + `*`）。行内代码单独一档，命中后不再解析里面的星号。
_INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")

_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET = re.compile(r"^\s*[-*]\s+(.*)$")
_ORDERED = re.compile(r"^\s*\d+[.)]\s+(.*)$")
_HRULE = re.compile(r"^\s*([-*_])\s*(\1\s*){2,}$")


def _die(msg: str) -> None:
    """统一的中文报错出口，措辞与 doc_text.py / pdf_tables.py 等脚本对齐。"""
    print(f"[doc-convert] 错误：{msg}", file=sys.stderr)
    raise SystemExit(2)


def _read_md(src: Path) -> str:
    """把 Markdown 读成文本，兼容 BOM 与 GBK。

    评审实测两个坑，都发生在"读文件"这一步，一个静默一个崩：

    1. **BOM**（字节顺序标记）：Windows 记事本、以及 VS Code 的部分默认配置，
       存 UTF-8 时会在文件最开头塞三个不可见字节。原来的
       `read_text(encoding="utf-8")` 会把它当成正文第一个字符保留下来，于是
       首行变成 `\\ufeff# 标题`——匹配不上 _HEADING，**被当成普通正文默默写进
       Word**。用户拿到的文档少了一级标题，脚本却 exit 0 说"已生成"，属于本
       分支最不能接受的那类"看起来正常实则有缺陷"。改用 utf-8-sig：它只在
       BOM 存在时吃掉它，对没有 BOM 的文件行为完全不变。
    2. **GBK**：中文 Windows 上另存的 .md 常是 GBK/GB2312 编码，
       原来会直接抛 UnicodeDecodeError 一屏英文堆栈。这里退一步用 gb18030
       重试（GB18030 是 GBK 的超集，能覆盖 GBK/GB2312 的全部字节）。
       顺序不能反：先严后宽，UTF-8 解不通才认为是国标编码，否则 UTF-8 中文
       会被 gb18030 解成乱码而不报错。
    """
    try:
        raw = src.read_bytes()
    except OSError as e:
        _die(f"读取「{src.name}」失败：{e}")
    for enc in ("utf-8-sig", "gb18030"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    _die(
        f"「{src.name}」的文字编码认不出来（既不是 UTF-8 也不是 GBK/GB18030）。"
        "请用记事本或 VS Code 把它另存为 UTF-8 编码后重试。"
    )
    return ""  # 不可达，只为类型完整


def _add_inline(paragraph, text: str) -> None:
    """把一行正文按行内标记切成多个 run 加进段落。"""
    for piece in _INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**") and len(piece) > 4:
            paragraph.add_run(piece[2:-2]).bold = True
        elif piece.startswith("`") and piece.endswith("`") and len(piece) > 2:
            run = paragraph.add_run(piece[1:-1])
            run.font.name = "Consolas"
        elif piece.startswith("*") and piece.endswith("*") and len(piece) > 2:
            paragraph.add_run(piece[1:-1]).italic = True
        else:
            paragraph.add_run(piece)


_SEP_CELL = re.compile(r":?-{3,}:?")


def _is_table_sep(line: str) -> bool:
    """判定分隔行（`|---|:---:|`）。它是「这一片竖线行是表格」的唯一凭证——
    没有它就维持旧行为当普通文本，不误伤正文里碰巧带竖线的行。"""
    s = line.strip()
    if not s.startswith("|"):
        return False
    parts = [p.strip() for p in s.strip("|").split("|")]
    return bool(parts) and all(_SEP_CELL.fullmatch(p) for p in parts)


def _split_cells(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _add_md_table(doc, header_line: str, sep_line: str,
                  data_lines: list[str], first_data_lineno: int) -> None:
    """把一片管道表格行画成真正的 Word 表格。

    边界纪律：短行补空单元格（GFM 标准行为，不丢信息）；长行拒绝——截断会
    丢内容，宁可报错让用户改（同本技能「不产出有缺陷文件」的头号纪律）。
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    headers = _split_cells(header_line)
    aligns = []
    for p in [c.strip() for c in sep_line.strip().strip("|").split("|")]:
        if p.startswith(":") and p.endswith(":"):
            aligns.append(WD_ALIGN_PARAGRAPH.CENTER)
        elif p.endswith(":"):
            aligns.append(WD_ALIGN_PARAGRAPH.RIGHT)
        else:
            aligns.append(None)  # 左对齐是 Word 默认，不用显式设
    # 分隔行列数与表头不齐时对齐信息按左对齐补齐——对齐是装饰，不值得拒绝
    aligns += [None] * (len(headers) - len(aligns))

    rows: list[list[str]] = []
    for offset, ln in enumerate(data_lines):
        cells = _split_cells(ln)
        if len(cells) > len(headers):
            _die(f"Markdown 表格第 {first_data_lineno + offset} 行有 "
                 f"{len(cells)} 列，多于表头的 {len(headers)} 列。多出来的列"
                 "如果截断就丢内容，请把这一行改成与表头一致的列数后重试。")
        cells += [""] * (len(headers) - len(cells))
        rows.append(cells)

    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"  # 带边框；无边框的表在 Word 里看不出是表
    for c, text in enumerate(headers):
        par = table.rows[0].cells[c].paragraphs[0]
        _add_inline(par, text)
        for run in par.runs:
            run.bold = True
        if aligns[c] is not None:
            par.alignment = aligns[c]
    for r, cells in enumerate(rows, start=1):
        for c, text in enumerate(cells):
            par = table.rows[r].cells[c].paragraphs[0]
            _add_inline(par, text)
            if aligns[c] is not None:
                par.alignment = aligns[c]


def convert(src: Path, dst: Path) -> None:
    """把 src 这份 Markdown 转成 dst 这份 .docx。"""
    doc = Document()
    in_code = False
    lines = _read_md(src).splitlines()
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        # 围栏代码块：进入后逐行原样成段，直到再遇到围栏。放在最前面判断——
        # 代码块里的 `# ` 是注释不是标题，任何结构规则（包括表格）都不该在
        # 这里生效。
        if line.strip().startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            run = doc.add_paragraph().add_run(line)
            run.font.name = "Consolas"
            i += 1
            continue

        # 管道表格：当前行以 | 开头、下一行是分隔行，才认作表格——分隔行是
        # 唯一凭证，没有它就落到下面当普通文本（不误伤正文里的竖线）。
        if (line.strip().startswith("|") and i + 1 < len(lines)
                and _is_table_sep(lines[i + 1])):
            j = i + 2
            while j < len(lines) and lines[j].strip().startswith("|"):
                j += 1
            # 行号从 1 起：数据行从文件的第 i+3 行开始（表头 i+1、分隔行 i+2）
            _add_md_table(doc, line, lines[i + 1],
                          [ln.rstrip() for ln in lines[i + 2:j]],
                          first_data_lineno=i + 3)
            i = j
            continue

        if not line.strip():
            i += 1
            continue

        if _HRULE.match(line):
            doc.add_paragraph("―" * 20)
            i += 1
            continue

        m = _HEADING.match(line)
        if m:
            doc.add_heading(m.group(2).strip(), level=len(m.group(1)))
            i += 1
            continue

        m = _BULLET.match(line)
        if m:
            _add_inline(doc.add_paragraph(style="List Bullet"), m.group(1))
            i += 1
            continue

        m = _ORDERED.match(line)
        if m:
            _add_inline(doc.add_paragraph(style="List Number"), m.group(1))
            i += 1
            continue

        _add_inline(doc.add_paragraph(), line)
        i += 1

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))


def main(argv: list[str] | None = None) -> int:
    try:
        ap = argparse.ArgumentParser(description="Markdown 转 Word")
        ap.add_argument("input", help="输入 .md 文件")
        ap.add_argument("-o", "--output", required=True, help="输出 .docx 文件")
        args = ap.parse_args(argv)

        src = Path(args.input)
        if not src.is_file():
            _die(f"找不到输入文件 {src}")

        dst = Path(args.output)
        convert(src, dst)
        print(f"[doc-convert] 已生成 {dst}")
        return 0
    except Exception as e:
        # 兜底：main() 必须有这一层。最现实的漏网是 doc.save() 遇到只读目录/
        # 磁盘满时抛的裸 OSError。同 doc_text.py / pdf_tables.py 的纪律——
        # 那两个脚本的注释里早就把本脚本列为"已有兜底"的同伴了，实际漏了。
        # SystemExit 继承 BaseException，主动报错退出不会被这里重新包一层。
        _die(f"处理过程中出错：{type(e).__name__}: {e}")
        return 2  # 不可达，只为类型完整


if __name__ == "__main__":
    raise SystemExit(main())
