import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))

import export


STYLE = {
    "body": "font-size:16px;",
    "h1": "font-size:22px;",
    "h2": "font-size:19px;",
    "h3": "font-size:17px;",
    "quote": "color:#666;",
    "strong": "font-weight:bold;",
    "em": "font-style:italic;",
    "li": "font-size:16px;",
    "hr": "border:none;",
}


def test_load_style_reads_bundled_preset():
    style = export.load_style("wechat-default")
    assert "body" in style
    assert "font-size" in style["body"]


def test_paragraph_gets_inline_style():
    html = export.md_to_wechat_html("这是一段正文。", STYLE)
    assert '<p style="font-size:16px;">这是一段正文。</p>' in html


def test_headings_mapped_to_levels():
    html = export.md_to_wechat_html("# 一级\n## 二级\n### 三级", STYLE)
    assert '<h1 style="font-size:22px;">一级</h1>' in html
    assert '<h2 style="font-size:19px;">二级</h2>' in html
    assert '<h3 style="font-size:17px;">三级</h3>' in html


def test_no_style_tag_emitted():
    # 公众号编辑器会剥掉 <style>，样式必须全内联
    html = export.md_to_wechat_html("# 标题\n正文", STYLE)
    assert "<style" not in html


def test_bold_and_italic_inline():
    html = export.md_to_wechat_html("这里**很重要**也*有点意思*。", STYLE)
    assert '<strong style="font-weight:bold;">很重要</strong>' in html
    assert '<em style="font-style:italic;">有点意思</em>' in html


def test_blockquote_and_list():
    html = export.md_to_wechat_html("> 引用一句\n\n- 第一条\n- 第二条", STYLE)
    assert '<blockquote style="color:#666;">引用一句</blockquote>' in html
    assert '<li style="font-size:16px;">第一条</li>' in html
    assert "<ul" in html


def test_html_escaped_in_body_text():
    html = export.md_to_wechat_html("a < b & c > d", STYLE)
    assert "&lt;" in html and "&amp;" in html and "&gt;" in html


def test_md_to_plain_strips_markup():
    plain = export.md_to_plain("# 标题\n\n这里**很重要**。\n\n- 一条")
    assert "#" not in plain
    assert "**" not in plain
    assert "很重要" in plain
    assert "一条" in plain


def test_asterisk_as_multiplication_not_italicized():
    # 「长*宽*高」里的星号是乘号 —— 收尾星号后紧跟汉字，不能被当斜体吞掉
    plain = export.md_to_plain("面积是 长*宽*高 的乘积")
    assert "长*宽*高" in plain
    html = export.md_to_wechat_html("面积是 长*宽*高 的乘积", STYLE)
    assert "<em" not in html


def test_real_italic_still_rendered():
    # 收尾星号后跟标点/结尾的，仍是正常斜体（现有用法不能被误伤）
    html = export.md_to_wechat_html("这里*有点意思*。", STYLE)
    assert '<em style="font-style:italic;">有点意思</em>' in html


def test_h4_to_h6_headings_not_leaked():
    # 4–6 级标题不能把 #### 原样漏进读者可见输出
    html = export.md_to_wechat_html("#### 四级标题", STYLE)
    assert "####" not in html
    assert "四级标题" in html
    plain = export.md_to_plain("##### 五级标题")
    assert "#" not in plain
    assert "五级标题" in plain


def test_parse_images_extracts_caption_src_and_line():
    md = "开头。\n\n![深夜的便利店](../images/gen-1.png)\n\n结尾。"
    refs = export.parse_images(md)
    assert len(refs) == 1
    assert refs[0].caption == "深夜的便利店"
    assert refs[0].src == "../images/gen-1.png"
    assert refs[0].line == 3


def test_parse_images_ignores_normal_links():
    assert export.parse_images("详见[报告](https://example.com/a)。") == []


def test_parse_images_skips_fenced_blocks():
    """mermaid / 代码块里出现的图片语法是示例文本，不是真配图。
    当成真配图会导致导出闸报「缺图」，卡住一次本该成功的导出。"""
    md = "正文。\n\n```markdown\n![示例](../images/nope.png)\n```\n"
    assert export.parse_images(md) == []


def test_resolve_image_path_is_relative_to_markdown_file(tmp_path):
    """正文在 drafts/、图在 images/，相对路径必须按 md 文件所在目录解析，
    不是按当前工作目录——否则从别处跑导出脚本就全找不到图。"""
    md_path = tmp_path / "drafts" / "01.md"
    resolved = export.resolve_image_path("../images/a.png", md_path)
    assert resolved == tmp_path / "images" / "a.png"


def test_resolve_image_path_passes_absolute_through(tmp_path):
    abs_src = str(tmp_path / "images" / "a.png")
    assert export.resolve_image_path(abs_src, tmp_path / "drafts" / "01.md") == Path(abs_src)


def test_missing_images_reports_only_absent_files(tmp_path):
    (tmp_path / "images").mkdir()
    (tmp_path / "images" / "there.png").write_bytes(b"x")
    (tmp_path / "drafts").mkdir()
    md_path = tmp_path / "drafts" / "01.md"
    md = "![在的](../images/there.png)\n\n![不在的](../images/gone.png)"
    md_path.write_text(md, encoding="utf-8")
    missing = export.missing_images(md, md_path)
    assert [r.src for r in missing] == ["../images/gone.png"]


def test_main_blocks_export_when_image_missing(tmp_path, capsys):
    """缺图必须停下报清单，而不是导出一份引用损坏的稿。
    下游（公众号编辑器 / Word）不会在这一层报错，带着缺口跑完
    只会产出一份看着成功、打开全是碎图的成品。"""
    (tmp_path / "drafts").mkdir()
    md_path = tmp_path / "drafts" / "01.md"
    md_path.write_text("正文。\n\n![缺的](../images/gone.png)\n", encoding="utf-8")
    code = export.main([str(md_path), "--format", "plain", "--out", str(tmp_path / "o.txt")])
    assert code == 1
    assert "gone.png" in capsys.readouterr().out
    assert not (tmp_path / "o.txt").exists()


def test_wechat_html_renders_image_with_caption():
    """公众号里图和图说是一体的：<img> 后跟一行居中小字图说。
    图说为空时不产出空的说明行（留着是一条视觉上莫名其妙的空隙）。"""
    style = export.load_style("wechat-default")
    html_out = export.md_to_wechat_html("![深夜的便利店](../images/gen-1.png)", style)
    assert 'src="../images/gen-1.png"' in html_out
    assert "深夜的便利店" in html_out
    assert "<p" not in html_out.split("<img")[0]  # 图不该被包成普通段落


def test_wechat_html_image_without_caption_has_no_caption_line():
    style = export.load_style("wechat-default")
    html_out = export.md_to_wechat_html("![](../images/gen-1.png)", style)
    assert "<img" in html_out
    assert "figcaption" not in html_out


def test_wechat_html_escapes_caption():
    """图说来自 AI 生成的文本，可能含 < >，不转义就把 HTML 结构打坏了。"""
    style = export.load_style("wechat-default")
    html_out = export.md_to_wechat_html('![a<b>c](../images/x.png)', style)
    assert "<b>" not in html_out
    assert "&lt;b&gt;" in html_out


def test_plain_export_renders_image_as_caption_marker():
    """纯文本没法放图，退化成一个人能看懂的占位标记，
    而不是把 markdown 语法原样漏给读者。"""
    out = export.md_to_plain("![深夜的便利店](../images/gen-1.png)")
    assert out == "［图：深夜的便利店］"


def test_docx_embeds_existing_image(tmp_path):
    import base64

    from docx import Document

    (tmp_path / "images").mkdir()
    (tmp_path / "drafts").mkdir()
    # 1x1 透明 PNG。用 base64 而不是手打 hex：hex 串抄错一个字符，
    # python-docx 报的是「无法识别的图片格式」，会被误当成实现有 bug。
    png = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNkYPhfDwAChwGA60e6kgAAAABJRU5ErkJggg=="
    )
    (tmp_path / "images" / "a.png").write_bytes(png)
    md_path = tmp_path / "drafts" / "01.md"
    out = tmp_path / "o.docx"
    export.md_to_docx("正文。\n\n![图说](../images/a.png)\n", out, md_path)
    doc = Document(str(out))
    assert len(doc.inline_shapes) == 1


def test_copy_images_numbers_files_in_document_order(tmp_path):
    """文件名带序号，是为了让人在公众号编辑器里能按顺序对着插——
    原始文件名（gen-1754…png）对人没有任何顺序信息。"""
    (tmp_path / "images").mkdir()
    (tmp_path / "drafts").mkdir()
    for name in ("b.png", "a.png"):
        (tmp_path / "images" / name).write_bytes(b"x")
    md_path = tmp_path / "drafts" / "01.md"
    refs = export.parse_images("![二](../images/b.png)\n\n![一](../images/a.png)")
    out_dir = tmp_path / "output"
    pairs = export.copy_images(refs, md_path, out_dir)
    assert [name for _, name in pairs] == ["01-b.png", "02-a.png"]
    assert (out_dir / "images" / "01-b.png").is_file()
    assert (out_dir / "images" / "02-a.png").is_file()


def test_build_image_manifest_marks_cover_separately():
    """公众号封面在编辑器里是独立上传项、不进正文。
    不单独标出来，用户会把封面当成正文第一张图插进去。"""
    refs = export.parse_images("![封面](../images/a.png)\n\n![流程](../images/b.png)")
    pairs = [(refs[0], "01-a.png"), (refs[1], "02-b.png")]
    text = export.build_image_manifest(pairs, cover_first=True)
    assert "封面" in text
    assert "01-a.png" in text and "02-b.png" in text
    assert "output/images/" in text


def test_build_image_manifest_without_cover_lists_all_inline():
    refs = export.parse_images("![流程](../images/b.png)")
    text = export.build_image_manifest([(refs[0], "01-b.png")], cover_first=False)
    assert "封面" not in text


def test_build_image_manifest_is_empty_without_images():
    assert export.build_image_manifest([], cover_first=False) == ""


# ---------------------------------------------------------------- 未出图的围栏块

GENIMAGE_MD = """开头一段。

```genimage
图说: 凌晨三点的便利店
夜晚的便利店内景，暖黄色顶灯，落地玻璃窗外在下雨。
平视视角，不要出现任何文字。
```

结尾一段。
"""

MERMAID_MD = """开头一段。

```mermaid
graph TD
  A[收到需求] --> B{一句话说得清吗}
```

结尾一段。
"""


def test_wechat_html_never_leaks_genimage_source():
    """P1a 阶段正文里全是没出图的 ```genimage 块。没有围栏处理时，
    每一行出图指令都会被当成一个 <p> 渲染进成品——用户拿到的是一份
    夹着 AI 提示词的公众号稿。"""
    html_out = export.md_to_wechat_html(GENIMAGE_MD, STYLE)
    assert "```" not in html_out
    assert "夜晚的便利店内景" not in html_out
    assert "平视视角" not in html_out


def test_wechat_html_renders_genimage_as_visible_placeholder():
    """不能只是删掉——删了用户就不知道这里本该有张图。
    渲染成一个看得见的占位框，带上图说，人一眼知道图该插在哪。"""
    html_out = export.md_to_wechat_html(GENIMAGE_MD, STYLE)
    assert "待出图" in html_out
    assert "凌晨三点的便利店" in html_out
    assert "border:" in html_out  # 样式内联，同 build_image_manifest 的做法


def test_genimage_caption_accepts_fullwidth_colon():
    """写手在中文正文里打全角冒号是本能，只认半角必然漏掉一半图说。"""
    md = "```genimage\n图说：全角冒号也要认\n画面描述。\n```"
    assert "全角冒号也要认" in export.md_to_wechat_html(md, STYLE)


def test_wechat_html_never_leaks_mermaid_source():
    """公众号渲染不了 mermaid，源码漏进正文就是一堆读者看不懂的代码。"""
    html_out = export.md_to_wechat_html(MERMAID_MD, STYLE)
    assert "graph TD" not in html_out
    assert "```" not in html_out


def test_wechat_html_mermaid_placeholder_says_source_is_kept():
    """mermaid 的源码留在 Markdown 原稿里，占位块要把这件事说清楚，
    否则用户以为图被导出弄丢了。"""
    html_out = export.md_to_wechat_html(MERMAID_MD, STYLE)
    assert "mermaid" in html_out
    assert "原稿" in html_out


def test_plain_collapses_fenced_blocks_to_one_line_markers():
    """纯文本（朋友圈/私域话术）会被直接复制粘贴出去，
    围栏源码漏出去是最糟的结果。退化成一行占位标记，同 ［图：…］ 的路子。"""
    out = export.md_to_plain(GENIMAGE_MD)
    assert "```" not in out
    assert "夜晚的便利店内景" not in out
    assert "［待出图：凌晨三点的便利店］" in out

    out2 = export.md_to_plain(MERMAID_MD)
    assert "graph TD" not in out2
    assert "［信息图" in out2


def test_docx_does_not_emit_fence_source(tmp_path):
    from docx import Document

    out = tmp_path / "o.docx"
    export.md_to_docx(GENIMAGE_MD + "\n" + MERMAID_MD, out, tmp_path / "01.md")
    texts = [p.text for p in Document(str(out)).paragraphs]
    joined = "\n".join(texts)
    assert "```" not in joined
    assert "夜晚的便利店内景" not in joined
    assert "graph TD" not in joined
    assert any("待出图" in t for t in texts)


def test_pending_genimage_blocks_counted():
    assert export.count_pending_genimage(GENIMAGE_MD) == 1
    assert export.count_pending_genimage(GENIMAGE_MD + GENIMAGE_MD) == 2
    assert export.count_pending_genimage(MERMAID_MD) == 0


def test_main_reports_pending_genimage_blocks(tmp_path, capsys):
    """P1a 的既定终态就是「一批 genimage 描述块（未出图），导出时提示用户自行出图」。
    不提示的话，用户看到的只是几个占位框，不知道还差几张、该干什么。"""
    md_path = tmp_path / "01.md"
    md_path.write_text(GENIMAGE_MD + GENIMAGE_MD, encoding="utf-8")
    code = export.main([str(md_path), "--format", "plain", "--out", str(tmp_path / "o.txt")])
    assert code == 0
    printed = capsys.readouterr().out
    assert "没出图" in printed
    assert "2" in printed  # 张数要报出来，用户才知道还差几张


def test_main_silent_when_nothing_pending(tmp_path, capsys):
    md_path = tmp_path / "01.md"
    md_path.write_text("只有正文，没有出图指令。", encoding="utf-8")
    export.main([str(md_path), "--format", "plain", "--out", str(tmp_path / "o.txt")])
    assert "没出图" not in capsys.readouterr().out


# ---------------------------------------------------------------- 复制后的图要被引用


def test_wechat_html_img_src_points_to_copied_file():
    """copy_images 把图复制进 <out_dir>/images/ 是为了让导出成一个自带图的包；
    HTML 却指回项目树里的原始相对路径，换个 --out 目录就全断（`--out ~/桌面/稿.html`
    时 ../images/x.png 解析到 ~/images/x.png）。"""
    refs = export.parse_images("![封面](../images/a.png)")
    html_out = export.md_to_wechat_html(
        "![封面](../images/a.png)", STYLE, image_names={refs[0].line: "01-a.png"}
    )
    assert 'src="images/01-a.png"' in html_out
    assert "../images/a.png" not in html_out


def test_wechat_html_without_mapping_keeps_original_src():
    """不传映射时保持原样——直接调用（预览、测试）不该被迫先跑复制。"""
    html_out = export.md_to_wechat_html("![封面](../images/a.png)", STYLE)
    assert 'src="../images/a.png"' in html_out


def test_main_wechat_html_references_copied_images(tmp_path):
    (tmp_path / "images").mkdir()
    (tmp_path / "images" / "a.png").write_bytes(b"x")
    (tmp_path / "drafts").mkdir()
    md_path = tmp_path / "drafts" / "01.md"
    md_path.write_text("正文。\n\n![封面](../images/a.png)\n", encoding="utf-8")
    out = tmp_path / "output" / "o.html"
    out.parent.mkdir()
    assert export.main([str(md_path), "--format", "wechat", "--out", str(out)]) == 0
    html_out = out.read_text(encoding="utf-8")
    assert 'src="images/01-a.png"' in html_out
    assert (tmp_path / "output" / "images" / "01-a.png").is_file()


# ---------------------------------------------------------------- 配图必须独占一行


def test_inline_images_flags_image_not_alone_on_line():
    md = "这句话里夹了一张![流程](../images/f.png)图。\n\n![独占一行的](../images/g.png)"
    bad = export.inline_images(md)
    assert [r.caption for r in bad] == ["流程"]
    assert bad[0].line == 1


def test_inline_images_ignores_fenced_examples():
    md = "```markdown\n正文里夹一张![示例](../images/x.png)图\n```"
    assert export.inline_images(md) == []


def test_main_blocks_export_when_image_not_standalone(tmp_path, capsys):
    """行内图会在三种格式里表现不一致：HTML/docx 把它原样当文字渲染出来，
    纯文本却正确替换成标记；同时它照样被复制、被编号，还可能被插图清单
    标成「封面（不要插进正文）」。与其在三处各修一遍，不如把「独占一行」
    钉成硬约定，在闸这里一次拦下。"""
    (tmp_path / "images").mkdir()
    (tmp_path / "images" / "f.png").write_bytes(b"x")
    (tmp_path / "drafts").mkdir()
    md_path = tmp_path / "drafts" / "01.md"
    md_path.write_text("这句话里夹了一张![流程](../images/f.png)图。\n", encoding="utf-8")
    out = tmp_path / "o.html"
    code = export.main([str(md_path), "--format", "wechat", "--out", str(out)])
    assert code == 1
    printed = capsys.readouterr().out
    assert "流程" in printed
    assert "第 1 行" in printed
    assert not out.exists()
