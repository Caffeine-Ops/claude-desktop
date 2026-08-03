#!/usr/bin/env python3
"""导出定稿到各平台格式。

用法：
    python3 scripts/export.py <md路径> --format wechat|plain|docx [--style wechat-default] [--out <路径>]

为什么自己写 Markdown → HTML 而不用现成库：公众号编辑器会剥掉
`<style>` 标签和 class，样式**必须全部内联**在每个元素的 style 属性上。
现成的 markdown 库输出的是干净的语义 HTML（靠外部样式表），粘进公众号
就是一片没有格式的黑字。这里的转换刻意只覆盖写作真正会用到的语法子集
（标题/段落/粗斜体/引用/列表/分隔线/图片），不追求 CommonMark 完备。

**有一份 TypeScript 孪生实现，改这里必须同步改那边**：
`apps/studio/electron/main/core/writingWechat.ts` 镜像了本文件的
`md_to_wechat_html`（桌面端要在渲染进程里实时预览公众号排版，跨不过
Python 边界，只能重写一份）。两边有测试互相钉对齐，但测试只覆盖已有
语法——**新增一种块级语法（图片、genimage/mermaid 占位块……）时，
测试不会自动发现那边缺了**，必须人肉同步过去。这条指针是反向的一半，
正向的一半写在 writingWechat.ts 的文件头。
"""

from __future__ import annotations

import argparse
import html
import json
import re
import shutil
import sys
from dataclasses import dataclass
from pathlib import Path

SKILL_DIR = Path(__file__).resolve().parent.parent
STYLES_DIR = SKILL_DIR / "templates" / "export_styles"

# #{1,6}：h1–h6 全收。只匹配 1–3 会让 #### 原样漏进读者可见输出，
# 且与 strip_markdown / readability 的 #{1,6} 口径不一致。
_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_QUOTE = re.compile(r"^>\s?(.*)$")
_LIST_ITEM = re.compile(r"^[-*]\s+(.*)$")
_HR = re.compile(r"^\s*(-{3,}|\*{3,})\s*$")
_BOLD = re.compile(r"\*\*(.+?)\*\*")
# 收尾星号后不能紧跟「词字」（汉字/字母/数字）——这一条把「长*宽*高」这种
# 用星号当乘号/脚注的写法排除掉（它的收尾 * 后跟着汉字），避免被当斜体吞掉星号；
# 正常斜体的收尾 * 后面是标点/空格/行尾（如「*有点意思*。」），照常识别。
# 内容不含 *（[^*]+?）以免跨过下一个星号误配。
_ITALIC = re.compile(r"(?<!\*)\*(?!\*)([^*]+?)\*(?![\w*])")

# 图片语法。前导 ! 是与普通链接的唯一区别；路径部分排除空白与右括号，
# 后面可选一个 markdown 标题后缀 `"..."`（`![图](路径 "标题")`）。
# 与 writing_utils._IMAGE_SYNTAX 是两份正则、口径刻意不同：那边只需要
# 「整体删掉」，这边要**分组取出** caption 与 src，合并成一份反而两头别扭。
_IMAGE = re.compile(r"!\[([^\]]*)\]\(([^)\s]+)(?:\s+\"[^\"]*\")?\)")

# 围栏块开头。lang 取 ``` 后面那个词（```genimage / ```mermaid），没写就是空串。
# 只用来识别下面 _PLACEHOLDER_LANGS 里的两种；其余围栏（```python 之类）
# 一行一行照旧走原路径，刻意不改——本次要修的是配图指令泄漏，不是给
# 这个只覆盖写作语法子集的转换器补一套通用代码块渲染。
_FENCE_OPEN = re.compile(r"^\s*```+\s*([^\s`]*)\s*$")

# 图说行。格式由 references/illustrator.md 规定：块的首行写 `图说: xxx`。
# 全角半角冒号都认——写手在中文正文里打全角冒号是本能，只认半角会漏掉一半图说，
# 而漏了图说的占位块等于只告诉用户「这儿有张图」却不说是哪张，帮助有限。
_CAPTION_LINE = re.compile(r"^\s*图说\s*[:：]\s*(.*?)\s*$")

# 这两种围栏块在 P1a 阶段都还没变成真图片：genimage 是待执行的出图指令，
# mermaid 是公众号渲染不了、需要人工出图再贴的信息图源码。两者都不能把
# 源码当正文段落漏进成品，也不能悄悄删掉（删了用户就不知道这里本该有图）。
_PLACEHOLDER_LANGS = ("genimage", "mermaid")


@dataclass
class FenceBlock:
    """一个待处理的围栏块（genimage / mermaid）。行号由 split_blocks 一并给出。"""

    lang: str
    lines: list[str]

    @property
    def caption(self) -> str:
        """块里的图说。取第一条匹配的 `图说:` 行。

        按 illustrator.md 的约定图说在首行，但这里扫全块而不是只看首行——
        写手偶尔会先写一行空行或把图说写在第二行，宽容一点不会误伤
        （正文描述行不会以「图说:」开头），严格一点却会白丢一条图说。
        """
        for raw in self.lines:
            m = _CAPTION_LINE.match(raw)
            if m:
                return m.group(1).strip()
        return ""


def split_blocks(markdown: str) -> list[tuple[int, str | FenceBlock]]:
    """把正文切成两种单元，各带 1-based 行号：普通行（str）与围栏块（FenceBlock）。

    三种导出格式都先过这一层，保证「genimage / mermaid 的源码绝不进成品」
    只实现一次——之前 HTML 和 docx 各自按行渲染，同一份稿在两种格式里
    漏出来的东西还不一样。行号一并带出，是因为 md_to_wechat_html 要靠它
    去查这一行的图被复制成了哪个文件名。

    不成对的围栏（开了没关）一路吃到文末：与 parse_images 的取舍同源，
    格式坏掉的围栏本来就会把全文渲染带歪，真去配对纠错比这层该担的事重得多。
    """
    units: list[tuple[int, str | FenceBlock]] = []
    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        m = _FENCE_OPEN.match(lines[i])
        if m and m.group(1).lower() in _PLACEHOLDER_LANGS:
            start = i + 1  # 1-based
            lang = m.group(1).lower()
            body: list[str] = []
            i += 1
            while i < len(lines) and not _FENCE_OPEN.match(lines[i]):
                body.append(lines[i])
                i += 1
            i += 1  # 跳过收尾的 ```
            units.append((start, FenceBlock(lang=lang, lines=body)))
            continue
        units.append((i + 1, lines[i]))
        i += 1
    return units


def count_pending_genimage(markdown: str) -> int:
    """还没出图的 genimage 块数量。导出末尾要据此提示用户「还差几张图」。"""
    return sum(
        1 for _, u in split_blocks(markdown) if isinstance(u, FenceBlock) and u.lang == "genimage"
    )


@dataclass
class ImageRef:
    """正文里的一处配图引用。line 是 1-based 行号——缺图报告要能让人直接跳过去。"""

    caption: str
    src: str
    line: int


def parse_images(markdown: str) -> list[ImageRef]:
    """抽出正文里的全部配图引用。

    **围栏代码块内的图片语法一律跳过**：那是示例文本（mermaid 块、
    教程里贴的 markdown 片段），不是真配图。当成真配图会让导出闸误报
    「缺图」，卡住一次本该成功的导出。
    """
    refs: list[ImageRef] = []
    in_fence = False
    for idx, raw in enumerate(markdown.splitlines(), start=1):
        if raw.lstrip().startswith("```"):
            # 单纯取反、不配对校验：一个孤立/不成对的 ``` 会让 in_fence
            # 从此再翻不回来，后面所有图片引用都被当成围栏内容漏检。
            # 接受这个代价——格式坏掉的围栏本来就会把全文渲染带歪，
            # 真去配对纠错（找下一个 ``` 收口、处理嵌套）比这道闸该担的事重得多。
            in_fence = not in_fence
            continue
        if in_fence:
            continue
        for m in _IMAGE.finditer(raw):
            refs.append(ImageRef(caption=m.group(1).strip(), src=m.group(2), line=idx))
    return refs


def resolve_image_path(src: str, md_path: Path) -> Path:
    """相对路径按**正文文件所在目录**解析，不是按 cwd。

    正文在 `<项目>/drafts/`、图在 `<项目>/images/`，相对路径恒为
    `../images/x.png`。若按 cwd 解析，从项目外任何地方跑导出都会找不到图，
    而且报的错是「文件不存在」而非「路径基准错了」，极难排查。
    """
    p = Path(src)
    return p if p.is_absolute() else (md_path.parent / p).resolve()


def missing_images(markdown: str, md_path: Path) -> list[ImageRef]:
    """返回磁盘上找不到的配图引用，保持正文中的出现顺序。"""
    return [r for r in parse_images(markdown) if not resolve_image_path(r.src, md_path).is_file()]


def inline_images(markdown: str) -> list[ImageRef]:
    """返回**没有独占一行**的配图引用（句子中间夹的图）。

    为什么这要当错误拦下、而不是想办法把它渲染出来：整条导出链都建立在
    「一行一图」上——HTML 与 docx 用 `_IMAGE.fullmatch(line)` 认图，
    copy_images 按出现顺序给图编号，插图清单按「第几行」告诉用户插在哪、
    并按位置认封面。夹在句子里的图会同时打坏这四件事，而且坏得不一致：
    纯文本正确地替换成了 ［图：…］ 标记，HTML 和 docx 却把 `![流程](...)`
    原样当文字渲染出来，同一份稿三种格式各说各话。与其在三处各修一遍
    （还得回答「图和它前后的半句话怎么排版」这种没有好答案的问题），
    不如把「配图必须独占一行」钉成硬约定，在闸这里一次拦下。
    约定的文字版写在 references/illustrator.md。
    """
    lines = markdown.splitlines()
    return [r for r in parse_images(markdown) if not _IMAGE.fullmatch(lines[r.line - 1].strip())]


def load_style(name: str) -> dict[str, str]:
    path = STYLES_DIR / f"{name}.json"
    data = json.loads(path.read_text(encoding="utf-8"))
    return {k: v for k, v in data.items() if k != "name"}


def _inline(text: str, style: dict[str, str]) -> str:
    """行内标记 → 内联样式的 HTML。先转义再替换，避免用户文本里的 < > 破坏结构。"""
    escaped = html.escape(text, quote=False)
    escaped = _BOLD.sub(lambda m: f'<strong style="{style["strong"]}">{m.group(1)}</strong>', escaped)
    escaped = _ITALIC.sub(lambda m: f'<em style="{style["em"]}">{m.group(1)}</em>', escaped)
    return escaped


def copy_images(refs: list[ImageRef], md_path: Path, out_dir: Path) -> list[tuple[ImageRef, str]]:
    """把正文用到的图复制进 `<out_dir>/images/`，按正文出现顺序编号。

    为什么要复制而不是让 HTML 指回项目里的原图：导出产物是**要交出去的一包**
    （发给自己手机、发给同事代发），指回项目路径的 HTML 换台机器就全断。
    为什么要重编号：原始文件名（gen-1754…png）对人零顺序信息，而公众号
    必须由人按顺序手工插图——序号就是那份操作顺序。
    """
    dest_dir = out_dir / "images"
    dest_dir.mkdir(parents=True, exist_ok=True)
    pairs: list[tuple[ImageRef, str]] = []
    for i, ref in enumerate(refs, start=1):
        source = resolve_image_path(ref.src, md_path)
        name = f"{i:02d}-{source.name}"
        shutil.copyfile(source, dest_dir / name)
        pairs.append((ref, name))
    return pairs


def build_image_manifest(pairs: list[tuple[ImageRef, str]], cover_first: bool) -> str:
    """贴在导出 HTML 顶部的插图说明块。

    为什么必须有这块东西：**微信编辑器会丢弃所有指向本地文件的图**
    （它只认已上传到微信服务器的图），也不支持 data URI。也就是说
    「粘一次全带图」在这个平台上做不到——这是平台限制，不是实现偷懒。
    能做的只有把手工步骤压到最低：图按序号命名、逐张列出该插在哪。
    样式内联，理由同全篇（公众号会剥掉 <style> 与 class）。
    """
    if not pairs:
        return ""
    rows: list[str] = []
    for idx, (ref, name) in enumerate(pairs):
        role = "封面（在编辑器的封面位单独上传，不要插进正文）" if (cover_first and idx == 0) else f"正文第 {ref.line} 行"
        cap = html.escape(ref.caption or "无图说", quote=False)
        rows.append(f"<li style=\"margin:0.3em 0;\">「{cap}」 → <code>output/images/{html.escape(name, quote=False)}</code>　·　{role}</li>")
    return (
        '<div style="border:1px dashed #cccccc;background:#fafafa;padding:1em 1.2em;margin:0 0 1.6em 0;'
        'font-size:14px;color:#666666;line-height:1.7;">'
        f'<strong style="color:#333333;">本文共 {len(pairs)} 张配图，需在公众号编辑器里手工插入</strong>'
        '<br />（微信会丢弃指向本地文件的图，这一步无法自动化）'
        f'<ol style="margin:0.6em 0 0 0;padding-left:1.4em;">{"".join(rows)}</ol>'
        '</div>'
    )


def fence_placeholder_html(block: FenceBlock) -> str:
    """未出图的围栏块 → 一个看得见的占位框。

    为什么是占位框而不是直接删掉：删了以后成品上什么都不剩，用户不知道
    这里本该有张图、更不知道图说是什么，等于把 P1a「先写指令、后出图」的
    半成品状态藏了起来。占位框把「这儿缺一张什么图」摆在原位。
    样式内联，理由同全篇（公众号编辑器会剥掉 <style> 与 class）。
    """
    cap = html.escape(block.caption or "（这个块没写图说）", quote=False)
    if block.lang == "genimage":
        title = "待出图"
        hint = "出图指令保留在 Markdown 原稿里；出图后把图片按 <code>![图说](路径)</code> 换掉这个块，再重跑导出"
        border, bg, fg = "#d9a441", "#fffaf0", "#8a6d3b"
    else:
        title = "待渲染的信息图（mermaid）"
        # 说清源码没丢，否则用户会以为导出把图弄丢了、回头去原稿里找不到而重画一遍。
        hint = "公众号编辑器不认 mermaid，源码原样保留在 Markdown 原稿里；请自行渲染成图片后插到这个位置"
        border, bg, fg = "#8ab4d9", "#f5f9fc", "#3f6b8f"
    return (
        f'<section style="border:1px dashed {border};background:{bg};padding:1em 1.2em;'
        'margin:1.4em 0;text-align:center;font-size:14px;line-height:1.7;'
        f'color:{fg};">'
        f'<strong style="color:{fg};">{title}</strong><br />'
        f'「{cap}」<br />'
        f'<span style="font-size:12px;opacity:0.85;">{hint}</span>'
        '</section>'
    )


def fence_placeholder_text(block: FenceBlock) -> str:
    """未出图的围栏块 → 一行纯文本占位标记。

    与图片的 ［图：…］ 标记同一套路子：纯文本会被直接复制粘贴出去
    （朋友圈/私域话术），漏出围栏源码是最糟的结果，但整块吞掉又会让人
    不知道少了什么，所以压成一行、保住图说。
    """
    if block.lang == "genimage":
        return f"［待出图：{block.caption or '未写图说'}］"
    return f"［信息图：{block.caption or 'mermaid 源码见 Markdown 原稿，需自行渲染'}］"


def md_to_wechat_html(
    markdown: str, style: dict[str, str], image_names: dict[int, str] | None = None
) -> str:
    """Markdown → 全内联样式的公众号 HTML。

    `image_names` 是 `{正文行号: 复制后的文件名}`，由 copy_images 的结果拼出来
    （见 main）。为什么 <img> 必须指向复制件而不是正文里的原始相对路径：
    导出产物是**要交出去的一包**，`--out ~/桌面/稿.html` 时复制件落在
    `~/桌面/images/01-x.png`，而原始路径 `../images/x.png` 会解析到
    `~/images/x.png`——每一张图都是断的。
    按**行号**索引而不是按 src：同一张图可能在正文里出现两次、要对应两个不同
    的复制件编号，src 当键会把后一次覆盖掉前一次。行号是一一对应的
    （「配图必须独占一行」由 inline_images 那道闸保证，一行至多一张图）。
    默认 None＝不改写，保持直接调用（预览、单测）时的两参数用法可用。
    """
    out: list[str] = []
    in_list = False
    image_names = image_names or {}

    def close_list() -> None:
        nonlocal in_list
        if in_list:
            out.append("</ul>")
            in_list = False

    for lineno, unit in split_blocks(markdown):
        if isinstance(unit, FenceBlock):
            close_list()
            out.append(fence_placeholder_html(unit))
            continue

        line = unit.rstrip()
        if not line.strip():
            close_list()
            continue

        # 图独占一行 → 图 + 图说，不包进 <p>（公众号里 <p> 的 margin 会把
        # 图和图说撑开成两块不相干的东西）。样式键用 .get 兜底：用户可能自带
        # 一份没有 img/figcaption 键的样式 JSON，KeyError 会让整次导出崩掉。
        m = _IMAGE.fullmatch(line.strip())
        if m:
            close_list()
            caption, src = m.group(1).strip(), m.group(2)
            # 有复制件就指向复制件（`images/NN-x.png`），理由见函数 docstring。
            copied = image_names.get(lineno)
            href = f"images/{copied}" if copied else src
            img_style = style.get("img", "display:block;max-width:100%;height:auto;margin:1.4em auto 0.4em auto;")
            out.append(f'<img src="{html.escape(href, quote=True)}" alt="{html.escape(caption, quote=True)}" style="{img_style}" />')
            if caption:
                cap_style = style.get("figcaption", "display:block;text-align:center;font-size:13px;color:#999999;")
                out.append(f'<figcaption style="{cap_style}">{html.escape(caption, quote=False)}</figcaption>')
            continue

        if _HR.match(line):
            close_list()
            out.append(f'<hr style="{style["hr"]}" />')
            continue

        m = _HEADING.match(line)
        if m:
            close_list()
            # 样式表只到 h3；4–6 级钳到 h3 渲染——公众号正文极少用到 h4+，
            # 用 h3 样式呈现远好过让 #### 泄漏，也不会 KeyError。
            level = min(len(m.group(1)), 3)
            tag = f"h{level}"
            out.append(f'<{tag} style="{style[tag]}">{_inline(m.group(2), style)}</{tag}>')
            continue

        m = _QUOTE.match(line)
        if m:
            close_list()
            out.append(f'<blockquote style="{style["quote"]}">{_inline(m.group(1), style)}</blockquote>')
            continue

        m = _LIST_ITEM.match(line)
        if m:
            if not in_list:
                out.append('<ul style="margin:1em 0;padding-left:1.4em;">')
                in_list = True
            out.append(f'<li style="{style["li"]}">{_inline(m.group(1), style)}</li>')
            continue

        close_list()
        out.append(f'<p style="{style["body"]}">{_inline(line, style)}</p>')

    close_list()
    return "\n".join(out)


def md_to_plain(markdown: str) -> str:
    """剥掉所有标记，只留可读文本。用于朋友圈/私域话术这类纯文本场景。"""
    lines: list[str] = []
    for _, unit in split_blocks(markdown):
        if isinstance(unit, FenceBlock):
            lines.append(fence_placeholder_text(unit))
            continue
        line = unit.strip()
        if _HR.match(line):
            continue
        # 纯文本没法放图，退化成人能看懂的占位标记——把 markdown 语法
        # 原样漏给读者（朋友圈/私域话术会被直接复制粘贴）是最糟的结果。
        line = _IMAGE.sub(lambda m: f"［图：{m.group(1).strip() or '配图'}］", line)
        line = _HEADING.sub(r"\2", line)
        line = _QUOTE.sub(r"\1", line)
        line = _LIST_ITEM.sub(r"· \1", line)
        line = _BOLD.sub(r"\1", line)
        line = _ITALIC.sub(r"\1", line)
        lines.append(line)
    # 折叠连续空行
    result: list[str] = []
    for line in lines:
        if not line and result and not result[-1]:
            continue
        result.append(line)
    return "\n".join(result).strip()


def md_to_docx(markdown: str, out_path: Path, md_path: Path) -> None:
    """导出 Word。依赖 python-docx（requirements.txt 已列）。

    多收一个 md_path：图片在正文里是相对路径（../images/x.png），
    要按**正文文件所在目录**解析才找得到——见 resolve_image_path 的注释。
    """
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        raise SystemExit("[writing] 错误：导出 docx 需要 python-docx，请先跑 bin/ensure-python.sh 装依赖")

    doc = Document()
    for _, unit in split_blocks(markdown):
        if isinstance(unit, FenceBlock):
            # 与纯文本用同一行标记：Word 稿也是给人读的，源码漏进去一样刺眼。
            doc.add_paragraph(fence_placeholder_text(unit))
            continue
        line = unit.strip()
        if not line or _HR.match(line):
            continue
        # 图独占一行 → 嵌图 + 图说段。宽度钉 5.5 英寸（A4 正文宽度），
        # 不钉会按图片像素尺寸铺开，大图直接溢出页面。
        m = _IMAGE.fullmatch(line)
        if m:
            caption, src = m.group(1).strip(), m.group(2)
            doc.add_picture(str(resolve_image_path(src, md_path)), width=Inches(5.5))
            if caption:
                doc.add_paragraph(caption, style="Caption")
            continue
        m = _HEADING.match(line)
        if m:
            doc.add_heading(_BOLD.sub(r"\1", m.group(2)), level=len(m.group(1)))
            continue
        m = _LIST_ITEM.match(line)
        if m:
            doc.add_paragraph(_BOLD.sub(r"\1", m.group(1)), style="List Bullet")
            continue
        m = _QUOTE.match(line)
        if m:
            doc.add_paragraph(_BOLD.sub(r"\1", m.group(1)), style="Intense Quote")
            continue
        doc.add_paragraph(_ITALIC.sub(r"\1", _BOLD.sub(r"\1", line)))
    doc.save(out_path)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="导出定稿")
    parser.add_argument("path")
    parser.add_argument("--format", choices=("wechat", "plain", "docx"), default="wechat")
    parser.add_argument("--style", default="wechat-default")
    parser.add_argument("--out", default=None)
    args = parser.parse_args(argv)

    src = Path(args.path)
    markdown = src.read_text(encoding="utf-8")

    # 图片就位闸：缺图停下报清单，绝不导出一份引用损坏的稿。
    # 下游导出器（公众号编辑器 / python-docx）不在这一层检测缺失，
    # 带着缺口跑完只会产出一份「看着成功、打开全是碎图」的成品。
    # 同源做法见 ppt-master 的 Image readiness GATE。
    missing = missing_images(markdown, src)
    if missing:
        print(f"[writing] ✗ 有 {len(missing)} 张配图找不到文件，导出已中止：")
        for ref in missing:
            print(f"  - 第 {ref.line} 行「{ref.caption or '无图说'}」→ {ref.src}")
        print("[writing] 请先把缺的图放到上述路径，或从正文里删掉这些引用，再重跑导出。")
        return 1

    # 独占一行闸：句子中间夹的图一律拦下。理由见 inline_images 的注释——
    # 三种格式对行内图的表现不一致（纯文本对、HTML/docx 把语法当文字渲染出来），
    # 而且它还会被复制、编号、甚至被插图清单标成「封面（不要插进正文）」。
    # 与缺图闸同形（报清单 + 退出码 1）：都是「人改一下再重跑」的格式问题，
    # 用同一种交互，用户不用学两套。
    inline = inline_images(markdown)
    if inline:
        print(f"[writing] ✗ 有 {len(inline)} 张配图没有独占一行，导出已中止：")
        for ref in inline:
            print(f"  - 第 {ref.line} 行「{ref.caption or '无图说'}」→ {ref.src}")
        print("[writing] 配图必须自己占一整行（前后各留一个空行），不能夹在句子中间。")
        print("[writing] 因为导出是按行认图、按行编号、按位置认封面的，夹在句子里三件事全会坏。")
        print("[writing] 请把这些图移到单独一行，再重跑导出。判据见 references/illustrator.md。")
        return 1

    suffix = {"wechat": ".html", "plain": ".txt", "docx": ".docx"}[args.format]
    out_path = Path(args.out) if args.out else src.with_suffix(suffix)

    if args.format == "wechat":
        style = load_style(args.style)
        refs = parse_images(markdown)
        # cover_first：有图就把第一张当封面。**刻意不去读契约的 image_plan**——
        # export.py 收的是一个 md 文件路径，不是项目路径，正文可能来自
        # `<cwd>/写作/` 这类没有契约的单文件场景，为此反推项目根既脆弱又多余。
        # 代价可控：判错时只是多提示一句「第一张是封面」，而漏提示会让用户
        # 把封面当正文图插错位——两种错的代价不对称，取宁可多提示的那边。
        cover_first = bool(refs)
        pairs = copy_images(refs, src, out_path.parent) if refs else []
        # 先复制、再渲染：<img> 要指向复制件的文件名，而文件名是复制这一步
        # 才定下来的（按正文顺序编号）。顺序反过来 HTML 就只能指回原始路径。
        body = md_to_wechat_html(markdown, style, {ref.line: name for ref, name in pairs})
        out_path.write_text(build_image_manifest(pairs, cover_first) + body, encoding="utf-8")
        if pairs:
            print(f"[writing] 已复制 {len(pairs)} 张配图到：{out_path.parent / 'images'}")
    elif args.format == "plain":
        out_path.write_text(md_to_plain(markdown), encoding="utf-8")
    else:
        md_to_docx(markdown, out_path, src)

    print(f"[writing] 已导出：{out_path}")

    # 未出图提示。P1a 的既定终态就是「正文里一批 genimage 描述块（还没出图），
    # 导出时提示用户自行出图」——占位框在成品里是看得见的，但用户不会去数，
    # 这一行把「还差几张」直接摆到终端上。不当错误处理（不返回 1）：
    # 带占位框的稿是这一阶段的**正常**产物，不是坏掉的产物。
    pending = count_pending_genimage(markdown)
    if pending:
        print(f"[writing] ⚠ 还有 {pending} 处氛围图没出图，已导出成占位框。")
        print("[writing] 出图后把图片按 ![图说](路径) 换掉对应的 ```genimage 块，再重跑一次导出。")
    return 0


if __name__ == "__main__":
    sys.exit(main())
